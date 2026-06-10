from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "dist"
OUT_PATH = OUT_DIR / "codex-cli-first-practice-print-manual.docx"
PDF_PATH = OUT_DIR / "codex-cli-first-practice-print-manual.pdf"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B677A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F5F7FA"
WARN_FILL = "FFF7D6"


def manual_steps() -> list[dict[str, object]]:
    return [
        {
            "num": 1,
            "title": "자유 주제 정하기",
            "time": "0:05-0:15",
            "purpose": "오늘 만들 웹 미니 프로젝트의 방향을 한 문장으로 정한다.",
            "facilitator": "주제는 자유입니다. 단, 오늘은 전체 서비스를 만들지 않고 버튼 또는 입력창 하나가 있는 작은 화면만 만듭니다.",
            "actions": [
                "종이나 메모장에 만들고 싶은 주제를 한 문장으로 쓴다.",
                "주제가 크면 오늘 만들 수 있는 작은 화면으로 줄인다.",
                "버튼 또는 입력창 중 어느 쪽이 자연스러운지 생각한다.",
            ],
            "prompt": None,
            "stuck": [
                "쇼핑몰은 상품 3개 소개 페이지로 줄인다.",
                "예약 시스템은 예약 문의 버튼을 누르면 안내 문구가 나오는 페이지로 줄인다.",
                "회원 기능, 결제, 로그인, 외부 API가 들어가면 오늘 범위에서 제외한다.",
            ],
        },
        {
            "num": 2,
            "title": "Node.js 확인과 설치",
            "time": "0:15-0:35",
            "purpose": "Codex CLI 설치와 로컬 웹 확인을 위해 Node.js와 npm이 있는지 확인한다.",
            "facilitator": "Node.js는 JavaScript를 내 컴퓨터에서 실행하게 해주는 프로그램입니다. npm은 Node.js와 같이 설치되는 도구 상자입니다.",
            "actions": [
                "Windows는 PowerShell, macOS는 Terminal을 연다.",
                "진행자 화면을 보고 Node.js와 npm 버전 확인을 따라 한다.",
                "버전이 나오면 설치를 건너뛴다.",
                "버전이 안 나오면 브라우저에서 구글을 열고 nodejs를 검색한다.",
                "공식 Node.js 사이트에 들어가 LTS라고 표시된 안정 버전을 설치한다.",
                "설치가 끝나면 터미널을 완전히 닫고 다시 연 뒤 버전을 다시 확인한다.",
            ],
            "prompt": None,
            "stuck": [
                "설치했는데도 버전이 안 나오면 터미널을 새로 연다.",
                "회사나 학교 컴퓨터라 설치가 막히면 진행자 화면을 보며 흐름을 따라간다.",
            ],
        },
        {
            "num": 3,
            "title": "Codex CLI 설치와 ChatGPT 로그인",
            "time": "0:35-0:50",
            "purpose": "터미널에서 Codex와 대화할 수 있는 상태를 만든다.",
            "facilitator": "오늘은 ChatGPT 계정으로 로그인합니다. API key는 쓰지 않습니다. 설치 명령은 수업 당일 공식 문서를 보고 진행자 화면을 따라 합니다.",
            "actions": [
                "진행자가 공식 Codex 설치 문서를 화면에 띄우는 것을 본다.",
                "진행자 화면의 설치 명령을 PowerShell 또는 Terminal에 붙여넣고 실행한다.",
                "설치가 끝나면 Codex 버전 확인을 진행자 화면대로 따라 한다.",
                "Codex를 실행한다.",
                "브라우저가 열리면 ChatGPT 계정으로 로그인하고 승인한다.",
                "터미널에 Codex 대화 화면이 보이면 성공이다.",
            ],
            "prompt": None,
            "stuck": [
                "오류가 나오면 지우지 말고 그대로 둔다.",
                "npm 오류가 나오면 터미널을 새로 열고 다시 시도한다.",
                "5분 이상 막히면 진행자 화면을 보며 Plan Mode 문장 작성부터 따라간다.",
            ],
        },
        {
            "num": 4,
            "title": "작업 폴더 만들기",
            "time": "0:50-0:55",
            "purpose": "Codex가 작업할 빈 폴더를 만든다.",
            "facilitator": "오늘 폴더 이름은 영어와 숫자만 씁니다. 한글이나 공백은 처음 실습에서 헷갈릴 수 있습니다.",
            "actions": [
                "바탕화면이나 문서 폴더 안에 새 폴더를 만든다.",
                "폴더 이름은 my-codex-project처럼 단순하게 만든다.",
                "터미널에서 그 폴더 위치로 이동한다.",
                "현재 터미널 위치가 방금 만든 폴더인지 확인한다.",
            ],
            "prompt": None,
            "stuck": [
                "폴더 이동이 어렵다면 진행자가 OS별로 직접 도와준다.",
                "폴더가 비어 있는 상태에서 시작하는 것이 가장 좋다.",
            ],
        },
        {
            "num": 5,
            "title": "Plan Mode로 계획 세우기",
            "time": "0:55-1:05",
            "purpose": "Codex에게 바로 만들라고 하지 않고 먼저 계획만 세우게 한다.",
            "facilitator": "지금부터는 만들지 않습니다. 먼저 계획만 세웁니다. 이 단계를 Plan Mode라고 부르겠습니다.",
            "actions": [
                "Codex에서 Plan Mode를 켤 수 있으면 진행자 안내에 따라 켠다.",
                "아래 요청문을 자기 주제에 맞게 바꿔 보낸다.",
                "Codex가 바로 파일을 만들려고 하면 즉시 멈추고 계획만 다시 요청한다.",
            ],
            "prompt": "Plan Mode로 진행하자. 아직 파일을 만들거나 수정하지 말고, 먼저 계획만 세워줘. 내 주제는 [내 주제]야. 파일은 README.md, index.html, style.css, script.js, test.js로 구성하고, 백엔드, 로그인, 데이터베이스, 결제, 외부 API는 쓰지 않는다. 버튼 또는 입력창 중 하나만 사용하고, Node.js 로컬 서버와 test.js로 확인할 수 있게 계획해줘.",
            "stuck": None,
        },
        {
            "num": 6,
            "title": "계획 검토와 수정 요청",
            "time": "1:05-1:15",
            "purpose": "사람이 계획을 읽고 한 번 수정 요청한 뒤 구현을 승인한다.",
            "facilitator": "Codex가 계획을 줬다고 바로 승인하지 않습니다. 오늘은 모두가 수정 요청을 한 번씩 합니다.",
            "actions": [
                "내 주제가 화면에 분명히 드러나는지 확인한다.",
                "파일 5개가 계획에 들어 있는지 확인한다.",
                "버튼 또는 입력창 동작이 하나만 있는지 확인한다.",
                "브라우저 확인과 테스트 방법이 포함되어 있는지 확인한다.",
                "기능이 크거나 어려우면 더 작게 줄여달라고 요청한다.",
                "계획이 적당하면 구현을 승인한다.",
            ],
            "prompt": "이 계획으로 진행하자. 이제 파일을 만들어줘. 구현이 끝나면 파일별로 무엇을 만들었는지 설명하고, 내가 브라우저에서 확인할 순서를 알려줘.",
            "stuck": None,
        },
        {
            "num": 7,
            "title": "Codex가 프로젝트 구현",
            "time": "1:15-1:30",
            "purpose": "Codex가 파일을 만들고 참가자는 파일 구성과 설명을 확인한다.",
            "facilitator": "지금은 참가자가 코드를 직접 고치지 않습니다. Codex가 어떤 파일을 만들고 어떤 명령을 실행하는지 봅니다.",
            "actions": [
                "Codex가 작업하는 동안 파일 생성 과정을 본다.",
                "구현이 끝나면 README.md, index.html, style.css, script.js, test.js가 있는지 확인한다.",
                "빠진 파일이 있으면 Codex에게 빠진 파일을 만들어달라고 요청한다.",
                "각 파일 역할을 한 문장씩 설명해달라고 요청한다.",
            ],
            "prompt": None,
            "stuck": None,
        },
        {
            "num": 8,
            "title": "브라우저에서 확인",
            "time": "1:30-1:40",
            "purpose": "VS Code 없이 Node.js 로컬 서버로 화면을 연다.",
            "facilitator": "오늘은 VS Code를 쓰지 않습니다. Node.js로 간단한 로컬 서버를 열어서 확인합니다.",
            "actions": [
                "Codex에게 Node.js로 로컬 서버를 실행하는 방법을 OS에 맞게 안내해달라고 요청한다.",
                "Codex가 알려준 순서대로 로컬 서버를 실행한다.",
                "터미널에 나온 주소를 브라우저에 입력한다.",
                "자기 주제, 버튼 또는 입력창, 화면 변화가 보이는지 확인한다.",
            ],
            "prompt": "이 프로젝트를 VS Code 없이 브라우저에서 확인하고 싶어. Node.js로 간단한 로컬 서버를 실행하는 방법을 내 OS에 맞게 안내해줘. 실행 후 어떤 주소를 브라우저에서 열어야 하는지도 알려줘.",
            "stuck": [
                "페이지가 안 열리면 터미널 메시지와 브라우저 메시지를 지우지 말고 Codex에게 보여준다.",
                "버튼이 안 되면 전체를 다시 만들지 말고 그 문제만 고쳐달라고 요청한다.",
            ],
        },
        {
            "num": 9,
            "title": "수동 테스트와 test.js 실행",
            "time": "1:40-1:52",
            "purpose": "눈으로 확인하고, Node.js 검사 파일도 실행한다.",
            "facilitator": "보이니까 끝이 아닙니다. 직접 눌러보고, test.js도 실행해 확인합니다.",
            "actions": [
                "페이지 제목과 내 주제가 보이는지 확인한다.",
                "버튼 또는 입력창이 보이는지 확인한다.",
                "버튼을 누르거나 입력했을 때 화면이 바뀌는지 확인한다.",
                "브라우저 창을 좁혀도 내용이 심하게 겹치지 않는지 확인한다.",
                "새로고침하면 기본 화면으로 돌아오는지 확인한다.",
                "Codex에게 test.js 실행 방법을 묻고 그대로 실행한다.",
                "실패 메시지가 나오면 그대로 Codex에게 보여주고 작은 수정만 요청한다.",
            ],
            "prompt": "test.js가 README.md, index.html, style.css, script.js 존재 여부, CSS/JS 연결 여부, 버튼 또는 입력창 존재 여부, README 설명과 실행 방법을 확인하는지 봐줘. 부족하면 test.js를 수정하고 실행 방법을 알려줘.",
            "stuck": None,
        },
        {
            "num": 10,
            "title": "GitHub에 올리기",
            "time": "1:52-2:00",
            "purpose": "Git 명령을 외우지 않고 Codex에게 저장소 연결, 커밋, push, 확인을 맡긴다.",
            "facilitator": "GitHub는 오늘 만든 프로젝트를 올려두는 저장소로만 사용합니다. Git 명령은 Codex에게 맡기되, 무엇을 하려는지는 확인합니다.",
            "actions": [
                "Codex에게 Git 설치 여부를 확인해달라고 요청한다.",
                "Git이 없으면 구글에서 git을 검색해 공식 사이트에서 설치한다.",
                "브라우저에서 GitHub에 로그인하거나 계정을 만든다.",
                "빈 저장소를 만들고 저장소 주소를 복사한다.",
                "Codex에게 현재 폴더를 저장소와 연결하고 첫 커밋과 push를 진행해달라고 요청한다.",
                "GitHub 웹페이지에서 파일 5개가 올라갔는지 확인한다.",
            ],
            "prompt": "이 폴더를 내가 만든 GitHub 저장소와 연결하고 싶어. 저장소 주소는 [주소]야. 먼저 현재 폴더와 파일 상태를 확인하고, 어떤 명령을 실행할지 설명해줘. 내가 확인하면 첫 커밋을 만들고 GitHub에 push한 뒤 정상 업로드됐는지 확인해줘.",
            "stuck": [
                "시간이 부족하면 저장소 생성까지만 하고 다음에 이어갈 요약을 Codex에게 요청한다.",
                "인증 화면이 열리면 브라우저에서 GitHub 로그인을 완료하고 터미널로 돌아온다.",
            ],
        },
    ]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D6DCE5", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def remove_paragraph_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def style_text(paragraph, text: str, *, bold=False, color=INK, size=10.5) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    remove_paragraph_border(p)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Codex CLI 첫 실습 인쇄 매뉴얼")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    style_text(
        sub,
        "준비물 없이 시작하는 2시간 따라하기 수업 | Node.js, Codex CLI, Plan Mode, 미니 프로젝트, 테스트, GitHub 업로드",
        color=MUTED,
        size=10.5,
    )


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def add_body(doc: Document, text: str, *, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_label:
        style_text(p, bold_label, bold=True)
        style_text(p, text)
    else:
        style_text(p, text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(INK)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(INK)


def add_callout(doc: Document, title: str, lines: list[str], fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="C9D2DE", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    style_text(p, title, bold=True, color=DARK_BLUE, size=10.5)
    for line in lines:
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        style_text(para, line, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths=(1.6, 4.9)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    for label, detail in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            set_cell_margins(cell)
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], LIGHT_BLUE)
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        style_text(p0, label, bold=True, color=DARK_BLUE, size=10)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        style_text(p1, detail, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_schedule(doc: Document) -> None:
    rows = [
        ("0:00-0:05", "수업 목표와 규칙 확인", "오늘 만들 범위 이해"),
        ("0:05-0:15", "자유 주제 정하기", "자기 프로젝트 방향 정하기"),
        ("0:15-0:35", "Node.js 확인과 설치", "Codex CLI 설치 준비"),
        ("0:35-0:50", "Codex CLI 설치와 로그인", "Codex 실행 성공"),
        ("0:50-1:05", "작업 폴더 만들기와 Plan Mode 시작", "바로 만들지 않고 먼저 계획"),
        ("1:05-1:15", "계획 검토와 수정 요청", "사람이 구조를 승인"),
        ("1:15-1:30", "Codex가 미니 프로젝트 구현", "파일 생성"),
        ("1:30-1:40", "Node 로컬 서버로 브라우저 확인", "화면 열기"),
        ("1:40-1:52", "수동 확인과 test.js 실행", "검증 체험"),
        ("1:52-2:00", "GitHub 저장소 연결과 업로드", "결과 공유"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    headers = ("시간", "할 일", "목표")
    widths = (1.05, 3.05, 2.4)
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        set_cell_width(cell, widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_text(p, text, bold=True, color=DARK_BLUE, size=9.5)
    set_repeat_table_header(table.rows[0])
    for time, task, goal in rows:
        cells = table.add_row().cells
        values = (time, task, goal)
        for idx, value in enumerate(values):
            set_cell_margins(cells[idx], top=70, bottom=70, start=110, end=110)
            set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            style_text(p, value, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_step(doc: Document, num: int, title: str, time: str, purpose: str, facilitator: str, actions: list[str], prompt: str | None = None, stuck: list[str] | None = None) -> None:
    add_h2(doc, f"{num}단계. {title} ({time})")
    add_body(doc, purpose, bold_label="목표: ")
    if facilitator:
        add_callout(doc, "진행자 멘트", [facilitator], fill=LIGHT_BLUE)
    if actions:
        add_body(doc, "참가자가 따라 할 행동")
        add_numbered(doc, actions)
    if prompt:
        add_callout(doc, "Codex에 보낼 문장", [prompt], fill="F7F9FC")
    if stuck:
        add_callout(doc, "막히면 이렇게 처리", stuck, fill=WARN_FILL)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_text(header, "Codex CLI 첫 실습 매뉴얼", color=MUTED, size=8.5)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_text(footer, "인쇄 배포용", color=MUTED, size=8.5)


def build_doc() -> Document:
    doc = Document()
    configure_doc(doc)
    add_title(doc)
    add_callout(
        doc,
        "오늘의 핵심",
        [
            "사람은 주제와 기준을 정하고, Codex는 계획, 구현, 수정, 검증, 업로드를 돕는다.",
            "코드를 직접 편집하지 않는다. 대신 Codex에게 정확히 요청하고 결과를 확인한다.",
            "최종 목표는 작은 정적 웹 프로젝트를 브라우저에서 확인하고 테스트까지 통과시키는 것이다.",
        ],
        fill=LIGHT_BLUE,
    )

    add_h1(doc, "1. 오늘 만들 결과물")
    add_body(doc, "각 참가자는 자기 주제로 작은 정적 웹 미니 프로젝트를 만든다. 공통 파일 구조는 다음과 같다.")
    add_kv_table(
        doc,
        [
            ("README.md", "프로젝트 목적, 실행 방법, 확인 방법을 설명한다."),
            ("index.html", "브라우저에 보이는 화면 구조를 담는다."),
            ("style.css", "색상, 간격, 글자 크기 같은 디자인을 담당한다."),
            ("script.js", "버튼 또는 입력창 동작을 담당한다."),
            ("test.js", "파일 존재, 연결 상태, 기본 구성 여부를 검사한다."),
        ],
    )

    add_h1(doc, "2. 전체 시간표")
    add_schedule(doc)

    add_h1(doc, "3. 수업 규칙")
    add_bullets(
        doc,
        [
            "사전 준비는 없다. 노트북과 인터넷만 있으면 시작한다.",
            "참가자는 코드를 직접 고치지 않는다. 파일 생성과 수정은 Codex에게 맡긴다.",
            "막히면 혼자 오래 붙잡지 않는다. 3분 이상 멈추면 진행자에게 알린다.",
            "백엔드, 로그인, 데이터베이스, 결제, 외부 API는 쓰지 않는다.",
            "개인정보, 비밀번호, 회사 내부 자료, API 키는 입력하지 않는다.",
            "명령어 암기가 아니라 계획, 구현, 확인, 검증, 업로드의 흐름을 익힌다.",
        ]
    )

    add_h1(doc, "4. 진행 단계")
    add_step(
        doc,
        1,
        "자유 주제 정하기",
        "0:05-0:15",
        "오늘 만들 웹 미니 프로젝트의 방향을 한 문장으로 정한다.",
        "주제는 자유입니다. 단, 오늘은 전체 서비스를 만들지 않고 버튼 또는 입력창 하나가 있는 작은 화면만 만듭니다.",
        [
            "종이나 메모장에 만들고 싶은 주제를 한 문장으로 쓴다.",
            "주제가 크면 오늘 만들 수 있는 작은 화면으로 줄인다.",
            "버튼 또는 입력창 중 어느 쪽이 자연스러운지 생각한다.",
        ],
        stuck=[
            "쇼핑몰은 상품 3개 소개 페이지로 줄인다.",
            "예약 시스템은 예약 문의 버튼을 누르면 안내 문구가 나오는 페이지로 줄인다.",
            "회원 기능, 결제, 로그인, 외부 API가 들어가면 오늘 범위에서 제외한다.",
        ],
    )

    add_step(
        doc,
        2,
        "Node.js 확인과 설치",
        "0:15-0:35",
        "Codex CLI 설치와 로컬 웹 확인을 위해 Node.js와 npm이 있는지 확인한다.",
        "Node.js는 JavaScript를 내 컴퓨터에서 실행하게 해주는 프로그램입니다. npm은 Node.js와 같이 설치되는 도구 상자입니다.",
        [
            "Windows는 PowerShell, macOS는 Terminal을 연다.",
            "진행자 화면을 보고 Node.js와 npm 버전 확인을 따라 한다.",
            "버전이 나오면 설치를 건너뛴다.",
            "버전이 안 나오면 브라우저에서 구글을 열고 nodejs를 검색한다.",
            "공식 Node.js 사이트에 들어가 LTS라고 표시된 안정 버전을 설치한다.",
            "설치가 끝나면 터미널을 완전히 닫고 다시 연 뒤 버전을 다시 확인한다.",
        ],
        stuck=[
            "설치했는데도 버전이 안 나오면 터미널을 새로 연다.",
            "회사나 학교 컴퓨터라 설치가 막히면 진행자 화면을 보며 흐름을 따라간다.",
        ],
    )

    add_step(
        doc,
        3,
        "Codex CLI 설치와 ChatGPT 로그인",
        "0:35-0:50",
        "터미널에서 Codex와 대화할 수 있는 상태를 만든다.",
        "오늘은 ChatGPT 계정으로 로그인합니다. API key는 쓰지 않습니다. 설치 명령은 수업 당일 공식 문서를 보고 진행자 화면을 따라 합니다.",
        [
            "진행자가 공식 Codex 설치 문서를 화면에 띄우는 것을 본다.",
            "진행자 화면의 설치 명령을 PowerShell 또는 Terminal에 붙여넣고 실행한다.",
            "설치가 끝나면 Codex 버전 확인을 진행자 화면대로 따라 한다.",
            "Codex를 실행한다.",
            "브라우저가 열리면 ChatGPT 계정으로 로그인하고 승인한다.",
            "터미널에 Codex 대화 화면이 보이면 성공이다.",
        ],
        stuck=[
            "오류가 나오면 지우지 말고 그대로 둔다.",
            "npm 오류가 나오면 터미널을 새로 열고 다시 시도한다.",
            "5분 이상 막히면 진행자 화면을 보며 Plan Mode 문장 작성부터 따라간다.",
        ],
    )

    add_step(
        doc,
        4,
        "작업 폴더 만들기",
        "0:50-0:55",
        "Codex가 작업할 빈 폴더를 만든다.",
        "오늘 폴더 이름은 영어와 숫자만 씁니다. 한글이나 공백은 처음 실습에서 헷갈릴 수 있습니다.",
        [
            "바탕화면이나 문서 폴더 안에 새 폴더를 만든다.",
            "폴더 이름은 my-codex-project처럼 단순하게 만든다.",
            "터미널에서 그 폴더 위치로 이동한다.",
            "현재 터미널 위치가 방금 만든 폴더인지 확인한다.",
        ],
        stuck=[
            "폴더 이동이 어렵다면 진행자가 OS별로 직접 도와준다.",
            "폴더가 비어 있는 상태에서 시작하는 것이 가장 좋다.",
        ],
    )

    add_step(
        doc,
        5,
        "Plan Mode로 계획 세우기",
        "0:55-1:05",
        "Codex에게 바로 만들라고 하지 않고 먼저 계획만 세우게 한다.",
        "지금부터는 만들지 않습니다. 먼저 계획만 세웁니다. 이 단계를 Plan Mode라고 부르겠습니다.",
        [
            "Codex에서 Plan Mode를 켤 수 있으면 진행자 안내에 따라 켠다.",
            "아래 요청문을 자기 주제에 맞게 바꿔 보낸다.",
            "Codex가 바로 파일을 만들려고 하면 즉시 멈추고 계획만 다시 요청한다.",
        ],
        prompt=(
            "Plan Mode로 진행하자. 아직 파일을 만들거나 수정하지 말고, 먼저 계획만 세워줘. "
            "내 주제는 [내 주제]야. 파일은 README.md, index.html, style.css, script.js, test.js로 구성하고, "
            "백엔드, 로그인, 데이터베이스, 결제, 외부 API는 쓰지 않는다. 버튼 또는 입력창 중 하나만 사용하고, "
            "Node.js 로컬 서버와 test.js로 확인할 수 있게 계획해줘."
        ),
    )

    add_step(
        doc,
        6,
        "계획 검토와 수정 요청",
        "1:05-1:15",
        "사람이 계획을 읽고 한 번 수정 요청한 뒤 구현을 승인한다.",
        "Codex가 계획을 줬다고 바로 승인하지 않습니다. 오늘은 모두가 수정 요청을 한 번씩 합니다.",
        [
            "내 주제가 화면에 분명히 드러나는지 확인한다.",
            "파일 5개가 계획에 들어 있는지 확인한다.",
            "버튼 또는 입력창 동작이 하나만 있는지 확인한다.",
            "브라우저 확인과 테스트 방법이 포함되어 있는지 확인한다.",
            "기능이 크거나 어려우면 더 작게 줄여달라고 요청한다.",
            "계획이 적당하면 구현을 승인한다.",
        ],
        prompt=(
            "이 계획으로 진행하자. 이제 파일을 만들어줘. 구현이 끝나면 파일별로 무엇을 만들었는지 설명하고, "
            "내가 브라우저에서 확인할 순서를 알려줘."
        ),
    )

    add_step(
        doc,
        7,
        "Codex가 프로젝트 구현",
        "1:15-1:30",
        "Codex가 파일을 만들고 참가자는 파일 구성과 설명을 확인한다.",
        "지금은 참가자가 코드를 직접 고치지 않습니다. Codex가 어떤 파일을 만들고 어떤 명령을 실행하는지 봅니다.",
        [
            "Codex가 작업하는 동안 파일 생성 과정을 본다.",
            "구현이 끝나면 README.md, index.html, style.css, script.js, test.js가 있는지 확인한다.",
            "빠진 파일이 있으면 Codex에게 빠진 파일을 만들어달라고 요청한다.",
            "각 파일 역할을 한 문장씩 설명해달라고 요청한다.",
        ],
    )

    add_step(
        doc,
        8,
        "브라우저에서 확인",
        "1:30-1:40",
        "VS Code 없이 Node.js 로컬 서버로 화면을 연다.",
        "오늘은 VS Code를 쓰지 않습니다. Node.js로 간단한 로컬 서버를 열어서 확인합니다.",
        [
            "Codex에게 Node.js로 로컬 서버를 실행하는 방법을 OS에 맞게 안내해달라고 요청한다.",
            "Codex가 알려준 순서대로 로컬 서버를 실행한다.",
            "터미널에 나온 주소를 브라우저에 입력한다.",
            "자기 주제, 버튼 또는 입력창, 화면 변화가 보이는지 확인한다.",
        ],
        prompt=(
            "이 프로젝트를 VS Code 없이 브라우저에서 확인하고 싶어. "
            "Node.js로 간단한 로컬 서버를 실행하는 방법을 내 OS에 맞게 안내해줘. "
            "실행 후 어떤 주소를 브라우저에서 열어야 하는지도 알려줘."
        ),
        stuck=[
            "페이지가 안 열리면 터미널 메시지와 브라우저 메시지를 지우지 말고 Codex에게 보여준다.",
            "버튼이 안 되면 전체를 다시 만들지 말고 그 문제만 고쳐달라고 요청한다.",
        ],
    )

    add_step(
        doc,
        9,
        "수동 테스트와 test.js 실행",
        "1:40-1:52",
        "눈으로 확인하고, Node.js 검사 파일도 실행한다.",
        "보이니까 끝이 아닙니다. 직접 눌러보고, test.js도 실행해 확인합니다.",
        [
            "페이지 제목과 내 주제가 보이는지 확인한다.",
            "버튼 또는 입력창이 보이는지 확인한다.",
            "버튼을 누르거나 입력했을 때 화면이 바뀌는지 확인한다.",
            "브라우저 창을 좁혀도 내용이 심하게 겹치지 않는지 확인한다.",
            "새로고침하면 기본 화면으로 돌아오는지 확인한다.",
            "Codex에게 test.js 실행 방법을 묻고 그대로 실행한다.",
            "실패 메시지가 나오면 그대로 Codex에게 보여주고 작은 수정만 요청한다.",
        ],
        prompt=(
            "test.js가 README.md, index.html, style.css, script.js 존재 여부, CSS/JS 연결 여부, "
            "버튼 또는 입력창 존재 여부, README 설명과 실행 방법을 확인하는지 봐줘. "
            "부족하면 test.js를 수정하고 실행 방법을 알려줘."
        ),
    )

    add_step(
        doc,
        10,
        "GitHub에 올리기",
        "1:52-2:00",
        "Git 명령을 외우지 않고 Codex에게 저장소 연결, 커밋, push, 확인을 맡긴다.",
        "GitHub는 오늘 만든 프로젝트를 올려두는 저장소로만 사용합니다. Git 명령은 Codex에게 맡기되, 무엇을 하려는지는 확인합니다.",
        [
            "Codex에게 Git 설치 여부를 확인해달라고 요청한다.",
            "Git이 없으면 구글에서 git을 검색해 공식 사이트에서 설치한다.",
            "브라우저에서 GitHub에 로그인하거나 계정을 만든다.",
            "빈 저장소를 만들고 저장소 주소를 복사한다.",
            "Codex에게 현재 폴더를 저장소와 연결하고 첫 커밋과 push를 진행해달라고 요청한다.",
            "GitHub 웹페이지에서 파일 5개가 올라갔는지 확인한다.",
        ],
        prompt=(
            "이 폴더를 내가 만든 GitHub 저장소와 연결하고 싶어. 저장소 주소는 [주소]야. "
            "먼저 현재 폴더와 파일 상태를 확인하고, 어떤 명령을 실행할지 설명해줘. "
            "내가 확인하면 첫 커밋을 만들고 GitHub에 push한 뒤 정상 업로드됐는지 확인해줘."
        ),
        stuck=[
            "시간이 부족하면 저장소 생성까지만 하고 다음에 이어갈 요약을 Codex에게 요청한다.",
            "인증 화면이 열리면 브라우저에서 GitHub 로그인을 완료하고 터미널로 돌아온다.",
        ],
    )

    add_h1(doc, "5. 마무리 회고")
    add_body(doc, "진행자는 마지막 1분에 아래 질문을 빠르게 확인한다.")
    add_bullets(
        doc,
        [
            "Codex에게 계획을 먼저 세우게 했는가?",
            "구현 전에 계획을 한 번 수정 요청했는가?",
            "브라우저에서 직접 확인했는가?",
            "test.js를 실행했는가?",
            "GitHub 업로드까지 갔는가, 아니면 어디서 멈췄는가?",
        ]
    )

    add_h1(doc, "6. 진행자 탈출 조건")
    add_bullets(
        doc,
        [
            "Node 설치가 막힌 사람은 진행자 화면을 보며 Plan Mode 문장만 작성한다.",
            "Codex 로그인 실패자는 진행자 화면을 따라 하며 주제와 계획 검토를 함께 한다.",
            "구현이 늦어지면 디자인 요구를 줄이고 버튼 또는 입력창 하나만 남긴다.",
            "브라우저 확인이 늦어지면 test.js보다 수동 확인을 우선한다.",
            "GitHub가 늦어지면 저장소 생성까지만 하고 다음 작업 요약을 만든다.",
        ]
    )

    return doc


def register_pdf_fonts() -> tuple[str, str]:
    font_candidates = [
        Path("C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/malgunbd.ttf"),
    ]
    regular = font_candidates[0]
    bold = font_candidates[1]
    if regular.exists():
        pdfmetrics.registerFont(TTFont("ManualKR", str(regular)))
        if bold.exists():
            pdfmetrics.registerFont(TTFont("ManualKRBold", str(bold)))
            return "ManualKR", "ManualKRBold"
        return "ManualKR", "ManualKR"
    return "Helvetica", "Helvetica-Bold"


def pdf_styles(font: str, bold_font: str):
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ManualTitle",
            parent=base["Title"],
            fontName=bold_font,
            fontSize=22,
            leading=28,
            textColor=colors.HexColor(f"#{DARK_BLUE}"),
            alignment=TA_CENTER,
            spaceAfter=10,
        ),
        "subtitle": ParagraphStyle(
            "ManualSubtitle",
            parent=base["Normal"],
            fontName=font,
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor(f"#{MUTED}"),
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
        "h1": ParagraphStyle(
            "ManualH1",
            parent=base["Heading1"],
            fontName=bold_font,
            fontSize=15,
            leading=19,
            textColor=colors.HexColor(f"#{BLUE}"),
            spaceBefore=14,
            spaceAfter=8,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "ManualH2",
            parent=base["Heading2"],
            fontName=bold_font,
            fontSize=12.2,
            leading=15,
            textColor=colors.HexColor(f"#{BLUE}"),
            spaceBefore=10,
            spaceAfter=6,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "ManualBody",
            parent=base["BodyText"],
            fontName=font,
            fontSize=9.3,
            leading=12.4,
            textColor=colors.HexColor(f"#{INK}"),
            spaceAfter=5,
        ),
        "small": ParagraphStyle(
            "ManualSmall",
            parent=base["BodyText"],
            fontName=font,
            fontSize=8.6,
            leading=11.2,
            textColor=colors.HexColor(f"#{INK}"),
            spaceAfter=3,
        ),
        "label": ParagraphStyle(
            "ManualLabel",
            parent=base["BodyText"],
            fontName=bold_font,
            fontSize=8.8,
            leading=11,
            textColor=colors.HexColor(f"#{DARK_BLUE}"),
            spaceAfter=2,
        ),
        "box": ParagraphStyle(
            "ManualBox",
            parent=base["BodyText"],
            fontName=font,
            fontSize=8.4,
            leading=11.2,
            textColor=colors.HexColor(f"#{INK}"),
            spaceAfter=2,
        ),
    }


def p(text: str, style) -> Paragraph:
    safe = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )
    return Paragraph(safe, style)


def pdf_bullets(items: list[str], style) -> ListFlowable:
    return ListFlowable(
        [ListItem(p(item, style), leftIndent=10) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=15,
        bulletFontSize=6,
        bulletOffsetY=1,
    )


def pdf_box(title: str, lines: list[str], styles, fill=LIGHT_GRAY) -> Table:
    data = [[p(title, styles["label"])], *[[p(line, styles["box"])] for line in lines]]
    table = Table(data, colWidths=[6.55 * inch], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{fill}")),
                ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D2DE")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def pdf_schedule(styles) -> Table:
    rows = [
        ("0:00-0:05", "수업 목표와 규칙 확인", "오늘 만들 범위 이해"),
        ("0:05-0:15", "자유 주제 정하기", "자기 프로젝트 방향 정하기"),
        ("0:15-0:35", "Node.js 확인과 설치", "Codex CLI 설치 준비"),
        ("0:35-0:50", "Codex CLI 설치와 로그인", "Codex 실행 성공"),
        ("0:50-1:05", "작업 폴더 만들기와 Plan Mode 시작", "먼저 계획"),
        ("1:05-1:15", "계획 검토와 수정 요청", "구조 승인"),
        ("1:15-1:30", "Codex가 미니 프로젝트 구현", "파일 생성"),
        ("1:30-1:40", "Node 로컬 서버로 브라우저 확인", "화면 열기"),
        ("1:40-1:52", "수동 확인과 test.js 실행", "검증 체험"),
        ("1:52-2:00", "GitHub 저장소 연결과 업로드", "결과 공유"),
    ]
    data = [[p("시간", styles["label"]), p("할 일", styles["label"]), p("목표", styles["label"])]]
    data.extend([[p(a, styles["small"]), p(b, styles["small"]), p(c, styles["small"])] for a, b, c in rows])
    table = Table(data, colWidths=[0.9 * inch, 3.15 * inch, 2.45 * inch], repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT_BLUE}")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D6DCE5")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def pdf_files_table(styles) -> Table:
    rows = [
        ("README.md", "프로젝트 목적, 실행 방법, 확인 방법을 설명한다."),
        ("index.html", "브라우저에 보이는 화면 구조를 담는다."),
        ("style.css", "색상, 간격, 글자 크기 같은 디자인을 담당한다."),
        ("script.js", "버튼 또는 입력창 동작을 담당한다."),
        ("test.js", "파일 존재, 연결 상태, 기본 구성 여부를 검사한다."),
    ]
    data = [[p(a, styles["label"]), p(b, styles["small"])] for a, b in rows]
    table = Table(data, colWidths=[1.25 * inch, 5.25 * inch], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{LIGHT_BLUE}")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D6DCE5")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def add_pdf_step(story: list, styles, step: dict[str, object]) -> None:
    title = f"{step['num']}단계. {step['title']} ({step['time']})"
    content = [
        p(title, styles["h2"]),
        p(f"목표: {step['purpose']}", styles["body"]),
        pdf_box("진행자 멘트", [str(step["facilitator"])], styles, fill=LIGHT_BLUE),
        Spacer(1, 5),
        p("참가자가 따라 할 행동", styles["label"]),
        pdf_bullets(list(step["actions"]), styles["small"]),
    ]
    if step.get("prompt"):
        content += [Spacer(1, 5), pdf_box("Codex에 보낼 문장", [str(step["prompt"])], styles, fill="F7F9FC")]
    if step.get("stuck"):
        content += [Spacer(1, 5), pdf_box("막히면 이렇게 처리", list(step["stuck"]), styles, fill=WARN_FILL)]
    content.append(Spacer(1, 8))
    story.append(KeepTogether(content))


def build_pdf() -> None:
    font, bold_font = register_pdf_fonts()
    styles = pdf_styles(font, bold_font)
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.72 * inch,
        leftMargin=0.72 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title="Codex CLI 첫 실습 인쇄 매뉴얼",
    )
    story = [
        p("Codex CLI 첫 실습 인쇄 매뉴얼", styles["title"]),
        p("준비물 없이 시작하는 2시간 따라하기 수업 | Node.js, Codex CLI, Plan Mode, 미니 프로젝트, 테스트, GitHub 업로드", styles["subtitle"]),
        pdf_box(
            "오늘의 핵심",
            [
                "사람은 주제와 기준을 정하고, Codex는 계획, 구현, 수정, 검증, 업로드를 돕는다.",
                "코드를 직접 편집하지 않는다. 대신 Codex에게 정확히 요청하고 결과를 확인한다.",
                "최종 목표는 작은 정적 웹 프로젝트를 브라우저에서 확인하고 테스트까지 통과시키는 것이다.",
            ],
            styles,
            fill=LIGHT_BLUE,
        ),
        Spacer(1, 10),
        p("1. 오늘 만들 결과물", styles["h1"]),
        p("각 참가자는 자기 주제로 작은 정적 웹 미니 프로젝트를 만든다. 공통 파일 구조는 다음과 같다.", styles["body"]),
        pdf_files_table(styles),
        Spacer(1, 10),
        p("2. 전체 시간표", styles["h1"]),
        pdf_schedule(styles),
        Spacer(1, 10),
        p("3. 수업 규칙", styles["h1"]),
        pdf_bullets(
            [
                "사전 준비는 없다. 노트북과 인터넷만 있으면 시작한다.",
                "참가자는 코드를 직접 고치지 않는다. 파일 생성과 수정은 Codex에게 맡긴다.",
                "막히면 혼자 오래 붙잡지 않는다. 3분 이상 멈추면 진행자에게 알린다.",
                "백엔드, 로그인, 데이터베이스, 결제, 외부 API는 쓰지 않는다.",
                "개인정보, 비밀번호, 회사 내부 자료, API 키는 입력하지 않는다.",
                "명령어 암기가 아니라 계획, 구현, 확인, 검증, 업로드의 흐름을 익힌다.",
            ],
            styles["small"],
        ),
        Spacer(1, 8),
        p("4. 진행 단계", styles["h1"]),
    ]
    for step in manual_steps():
        add_pdf_step(story, styles, step)
    story += [
        p("5. 마무리 회고", styles["h1"]),
        p("진행자는 마지막 1분에 아래 질문을 빠르게 확인한다.", styles["body"]),
        pdf_bullets(
            [
                "Codex에게 계획을 먼저 세우게 했는가?",
                "구현 전에 계획을 한 번 수정 요청했는가?",
                "브라우저에서 직접 확인했는가?",
                "test.js를 실행했는가?",
                "GitHub 업로드까지 갔는가, 아니면 어디서 멈췄는가?",
            ],
            styles["small"],
        ),
        p("6. 진행자 탈출 조건", styles["h1"]),
        pdf_bullets(
            [
                "Node 설치가 막힌 사람은 진행자 화면을 보며 Plan Mode 문장만 작성한다.",
                "Codex 로그인 실패자는 진행자 화면을 따라 하며 주제와 계획 검토를 함께 한다.",
                "구현이 늦어지면 디자인 요구를 줄이고 버튼 또는 입력창 하나만 남긴다.",
                "브라우저 확인이 늦어지면 test.js보다 수동 확인을 우선한다.",
                "GitHub가 늦어지면 저장소 생성까지만 하고 다음 작업 요약을 만든다.",
            ],
            styles["small"],
        ),
    ]
    doc.build(story)


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = build_doc()
    doc.save(OUT_PATH)
    build_pdf()
    print(OUT_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
